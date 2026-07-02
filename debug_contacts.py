#!/usr/bin/env python3
"""
Debug script to see what's in the Word document
"""

import docx

def debug_document():
    """Debug the Word document to see all content"""
    
    doc = docx.Document('Jesuit and Catholic High Schools in Seattle and Portland.docx')
    
    print("=== DEBUGGING WORD DOCUMENT ===")
    print()
    
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print()
    
    # Check all tables
    for i, table in enumerate(doc.tables):
        print(f"Table {i+1}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.columns)}")
        
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  Row {j+1}: {cells}")
        print()
    
    # Check paragraphs for any additional content
    print("Paragraphs:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  {i+1}: {para.text.strip()}")

if __name__ == "__main__":
    debug_document()
